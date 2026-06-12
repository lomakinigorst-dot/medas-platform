"""Генератор SITE_PLAN_v2.0.docx из SITE_PLAN.md"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Стили страницы
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

# Цвета MEDAS
PRIMARY = RGBColor(0, 48, 135)
GREEN = RGBColor(0, 169, 130)
DARK = RGBColor(25, 28, 30)
GRAY = RGBColor(67, 70, 85)
LIGHT_BG = RGBColor(247, 249, 251)

def set_font(run, name="Inter", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, "Manrope", 22, bold=True, color=PRIMARY)
    p.space_after = Pt(6)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "Manrope", 15, bold=True, color=PRIMARY)
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "Manrope", 12, bold=True, color=GREEN)
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    return p

def add_h4(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "Manrope", 11, bold=True, color=DARK)
    p.space_before = Pt(7)
    p.space_after = Pt(2)
    return p

def add_body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "Inter", 10, bold=bold, color=DARK)
    p.space_after = Pt(2)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, "Inter", 10, color=DARK)
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.space_after = Pt(1)
    return p

def add_divider():
    p = doc.add_paragraph()
    run = p.add_run("─" * 80)
    set_font(run, "Inter", 8, color=RGBColor(195, 198, 215))
    p.space_before = Pt(4)
    p.space_after = Pt(4)

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        set_font(run, "Inter", 9, bold=(bold_first and i == 0), color=DARK)
        p.space_after = Pt(1)

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # Шапка
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, "Inter", 9, bold=True, color=RGBColor(255, 255, 255))
        # Синий фон для шапки
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "003087")
        tcPr.append(shd)
    # Данные
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            # Очистить символы статусов для чистоты
            clean = str(text).replace("✅", "OK").replace("⛔", "404").replace("❌", "-").replace("🔴", "High").replace("🟡", "Mid").replace("🔵", "Low").replace("💡", "Idea")
            run = p.add_run(clean)
            set_font(run, "Inter", 8.5, color=DARK)
    return table

# ═══════════════════════════════════════════════════════════
# ТИТУЛЬНАЯ СТРАНИЦА
# ═══════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MEDAS")
set_font(run, "Manrope", 36, bold=True, color=PRIMARY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Структура и план сайта")
set_font(run, "Manrope", 20, bold=True, color=GREEN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SITE_PLAN v2.0 | Июнь 2026")
set_font(run, "Inter", 12, color=GRAY)

doc.add_paragraph()
add_divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("saas.med-as.ru  |  Синхронизировано с ТЗ v1.0")
set_font(run, "Inter", 10, color=GRAY)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. ДИЗАЙН-СИСТЕМА
# ═══════════════════════════════════════════════════════════
add_h2("1. Дизайн-система")

design_rows = [
    ("Primary", "#003087 (тёмно-синий)"),
    ("Secondary", "#00a982 (зелёный)"),
    ("Background", "#f7f9fb"),
    ("Text primary", "#191c1e"),
    ("Text secondary", "#434655"),
    ("Border", "#c3c6d7"),
    ("Success", "#e3fcef / #006644"),
    ("Danger", "#ba1a1a"),
    ("Шрифт заголовки", "Manrope (extrabold)"),
    ("Шрифт текст", "Inter"),
    ("Радиусы", "rounded-2xl (карточки), rounded-3xl (модалки), rounded-xl (кнопки)"),
    ("Тени", "shadow-xl shadow-[#003087]/8 (карты), shadow-sm (панели)"),
    ("Breakpoints", "lg = 1024px (мобиль / десктоп)"),
]
make_table(["Параметр", "Значение"], design_rows)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 2. ПОРЯДОК РАБОТЫ
# ═══════════════════════════════════════════════════════════
add_h2("2. Порядок работы — мастер-таблица")
add_body("Работаем постранично. Каждую страницу согласовываем до начала вёрстки.", bold=False)
doc.add_paragraph()

order_headers = ["№", "URL", "Название", "Stitch", "Статус", "Приоритет"]
order_rows = [
    ("1", "/", "Главная", "OK", "OK задеплоена", "High"),
    ("2", "/search", "Поиск врачей", "OK", "OK задеплоена", "High"),
    ("3", "/doctor/[slug]", "Профиль врача", "OK", "OK задеплоен", "High"),
    ("4", "/doctor/[slug]/booking", "Запись к врачу", "OK", "OK задеплоена", "High"),
    ("5", "/clinic/[slug]", "Профиль клиники", "OK", "OK задеплоен", "High"),
    ("6", "/services", "Услуги", "OK", "OK задеплоена", "Mid"),
    ("7", "/login", "Вход", "-", "OK задеплоена", "High"),
    ("8", "/register", "Регистрация", "-", "404 нет", "High"),
    ("9", "/forgot-password", "Восстановление пароля", "-", "404 нет", "Mid"),
    ("10", "/cabinet/patient", "Кабинет пациента", "OK", "OK задеплоен", "Mid"),
    ("11", "/cabinet/patient/appointments", "Мои записи", "-", "404 нет", "High"),
    ("12", "/cabinet/patient/medcard", "Медкарта", "OK", "OK задеплоена", "Mid"),
    ("13", "/cabinet/patient/family", "Семья", "OK", "OK задеплоена", "Mid"),
    ("14", "/cabinet/patient/bonuses", "Бонусы", "OK", "OK задеплоена", "Mid"),
    ("15", "/cabinet/patient/favorites", "Избранное", "-", "404 нет", "Mid"),
    ("16", "/cabinet/patient/settings", "Настройки пациента", "-", "404 нет", "Mid"),
    ("17", "/cabinet/clinic", "Кабинет клиники", "OK", "OK задеплоен", "Mid"),
    ("18", "/cabinet/clinic/doctors", "Врачи клиники", "OK", "404 нет", "Mid"),
    ("19", "/cabinet/clinic/schedule", "Расписание клиники", "-", "404 нет", "Mid"),
    ("20", "/cabinet/clinic/appointments", "Записи клиники", "-", "404 нет", "Mid"),
    ("21", "/cabinet/clinic/reports", "Отчёты", "OK", "OK задеплоена", "Mid"),
    ("22", "/cabinet/clinic/settings", "Настройки клиники", "-", "404 нет", "Mid"),
    ("23", "/cabinet/doctor", "Кабинет врача", "OK", "OK задеплоен", "Mid"),
    ("24", "/cabinet/doctor/schedule", "Расписание врача", "-", "404 нет", "Mid"),
    ("25", "/cabinet/doctor/appointments", "Приёмы врача", "-", "404 нет", "Mid"),
    ("26", "/cabinet/doctor/patients", "Пациенты", "-", "404 нет", "Mid"),
    ("27", "/cabinet/doctor/profile", "Профиль врача (ред.)", "-", "404 нет", "Mid"),
    ("28", "/services/[slug]", "Страница специальности", "-", "404 нет", "Mid"),
    ("29-35", "/articles, /about, /for-clinics...", "Блог + Контент + Юридика", "-", "404 нет", "Low"),
]
make_table(order_headers, order_rows)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 3. ПУБЛИЧНЫЕ СТРАНИЦЫ
# ═══════════════════════════════════════════════════════════
add_h2("3. Публичные страницы")

pages = [
    {
        "title": "/ — Главная страница",
        "status": "Задеплоена | Stitch: Страницы СТИЧ/Главная/ | Порядок: №1",
        "has": [
            "Header с навигацией + кнопки Войти / Зарегистрироваться",
            "Hero: заголовок, поисковая форма (специальность + город + дата)",
            "Блок специальностей (8 карточек с иконками)",
            "Блок «Как это работает» (4 шага)",
            "Топ-4 врача + Блок клиник (3 карточки) + Блок бонусов + Footer",
        ],
        "tz": [
            "Hero: поиск по симптому / имени (не только специальности)",
            "Блок «Доверие»: счётчики (500 000+ пациентов, 10 000+ врачей, 95% довольных)",
            "Блок «Акции» — карточки текущих акций клиник",
            "Блок «Статьи» — 3 последних медицинских статьи",
            "SEO: title, description, og-теги",
            "Мобильное гамбургер-меню",
        ],
        "delta": [
            "Блок счётчиков доверия (анимированные цифры)",
            "Секция «Акции» (2–3 карточки заглушки)",
            "Секция «Полезные статьи» (3 превью → /articles)",
            "SEO meta-теги в page.tsx",
            "Мобильное меню (гамбургер + drawer)",
        ],
    },
    {
        "title": "/search — Поиск врачей",
        "status": "Задеплоена | Stitch: Страницы СТИЧ/Врачи/ | Порядок: №2",
        "has": [
            "Сайдбар фильтров: специальность, доступность, стаж, цена, вызов на дом, пол",
            "3 карточки врачей (статика), пагинация",
        ],
        "tz": [
            "Строка поиска вверху (по имени / симптому / специальности)",
            "Фильтры: + район/метро, + язык, + ДМС, + онлайн",
            "Альтернативный вид — карта (Яндекс/Leaflet)",
            "Сортировка: рейтинг / цена / ближайшая доступность",
            "Поиск клиник: /search?type=clinic",
            "Чипы активных фильтров + сохранение в URL",
        ],
        "delta": [
            "Поисковая строка сверху",
            "Расширить фильтры (+метро, +ДМС, +онлайн)",
            "Чипы активных фильтров с крестиком",
            "Фильтры → URL-параметры (useState + useSearchParams)",
            "9 карточек вместо 3",
            "Переключатель «Список / Карта» (карта — заглушка)",
        ],
    },
    {
        "title": "/doctor/[slug] — Профиль врача",
        "status": "Задеплоен | Stitch: Страницы СТИЧ/Врач подробно/ | Порядок: №3",
        "has": [
            "Фото, имя, специальность, рейтинг, стаж, клиника",
            "Вкладки: О враче / Образование / Отзывы",
            "Sticky виджет записи (цена, мини-календарь, слоты)",
        ],
        "tz": [
            "Вкладка «Цены» — таблица услуг",
            "Вкладка «Клиника» — адрес, карта, режим работы",
            "Значки: верифицирован / ДМС / онлайн / выезд на дом",
            "Языки врача + Галерея (сертификаты)",
            "FAQ-аккордеон + SEO meta + schema.org Physician",
        ],
        "delta": [
            "Вкладка «Цены» с таблицей услуг",
            "Вкладка «Клиника» с реквизитами",
            "Значки (ДМС / онлайн / выезд)",
            "SEO meta + schema.org",
            "Галерея — 3–4 заглушки с иконкой документа",
            "FAQ-аккордеон (3–4 вопроса статически)",
        ],
    },
    {
        "title": "/doctor/[slug]/booking — Запись к врачу",
        "status": "Задеплоена | Stitch: Страницы СТИЧ/Запись к врачу с бонусами/ | Порядок: №4",
        "has": [
            "Мини-карточка врача, календарь, слоты, тип приёма, комментарий",
            "Блок бонусов (+450 баллов), итог записи",
        ],
        "tz": [
            "Поля ФИО + телефон + email (для незалогиненных)",
            "Выбор страховки (если принимает ДМС)",
            "Страница успеха /booking/success",
            "Шаг выбора члена семьи",
            "Применить бонусы (промокод / баллы)",
        ],
        "delta": [
            "Поля ФИО + телефон + email",
            "Блок «Применить бонусы»",
            "Шаг «Для кого запись»",
            "Страница /booking/success с деталями",
        ],
    },
    {
        "title": "/clinic/[slug] — Профиль клиники",
        "status": "Задеплоен | Stitch: Страницы СТИЧ/Профиль клиники/ | Порядок: №5",
        "has": [
            "Заголовок, рейтинг, адрес, режим работы",
            "О клинике (3 стата), 2 врача, галерея (заглушки), 9 услуг, 2 отзыва",
        ],
        "tz": [
            "Карта (Яндекс iframe)",
            "Галерея с лайтбоксом",
            "Все врачи клиники с пагинацией",
            "schema.org MedicalBusiness",
        ],
        "delta": [
            "Яндекс Карты embed (iframe)",
            "Галерея с иконками + лайтбокс JS",
            "6–8 карточек врачей + «Показать всех»",
            "schema.org MedicalBusiness",
        ],
    },
    {
        "title": "/services — Каталог услуг",
        "status": "Задеплоена | Stitch: Страницы СТИЧ/Услуги и диагностика/ | Порядок: №6",
        "has": ["Hero с поиском, 4 стат-блока, 9 специальностей, CTA"],
        "tz": [
            "Реальный поиск по услугам (JS-фильтрация)",
            "Фильтр: Диагностика / Лечение / Консультация",
            "Каждая карточка → /services/[slug]",
            "Блок «Популярные симптомы»",
        ],
        "delta": [
            "JS-фильтрация карточек по тексту",
            "Tabs: Все / Диагностика / Лечение",
            "Ссылки на /services/[slug]",
            "10 чипов симптомов",
        ],
    },
    {
        "title": "/login — Вход",
        "status": "Задеплоена | Stitch: нет | Порядок: №7",
        "has": ["Табы Пациент/Врач/Клиника, форма логин+пароль, SMS вход"],
        "tz": [
            "Основной флоу: телефон → SMS-код (6 цифр)",
            "Альтернатива: email + пароль",
            "Rate limiting UI (5 попыток → блокировка 10 мин)",
            "Госуслуги OAuth — заглушка «Скоро»",
        ],
        "delta": [
            "Форма: телефон + кнопка «Получить код»",
            "Второй шаг: 6-значный инпут SMS",
            "Ссылка «Войти по email»",
            "Табы реально переключают форму",
        ],
    },
    {
        "title": "/register — Регистрация",
        "status": "404 нет | Stitch: нет | Порядок: №8",
        "has": [],
        "tz": [
            "Шаг 1: выбор роли (Пациент / Врач / Клиника)",
            "Пациент (2 шага): телефон → SMS → ФИО + дата рождения + пол",
            "Врач (3 шага): телефон → SMS → ФИО + специальность + лицензия",
            "Клиника (3 шага): ИНН → ЕГРЮЛ → контакты → условия",
            "Соглашение на обработку мед.данных (обязательный чекбокс)",
            "+500 бонусных баллов при регистрации",
        ],
        "delta": ["Многошаговая форма с прогресс-баром (useState)"],
    },
]

for page in pages:
    add_h3(page["title"])
    add_body(page["status"], bold=False)
    if page["has"]:
        add_body("Что сейчас есть:", bold=True)
        for item in page["has"]:
            add_bullet(item)
    add_body("Требования ТЗ:", bold=True)
    for item in page["tz"]:
        add_bullet(item)
    add_body("Дельта (что делаем):", bold=True)
    for item in page["delta"]:
        add_bullet(item)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 4. ЛИЧНЫЙ КАБИНЕТ ПАЦИЕНТА
# ═══════════════════════════════════════════════════════════
add_h2("4. Личный кабинет пациента")

cabinet_patient_pages = [
    {
        "title": "/cabinet/patient — Дашборд пациента",
        "status": "Задеплоен | Stitch: ЛК пациента главная | Порядок: №10",
        "has": ["Sidebar (6 пунктов), баннер, запись, бонусы, показатели здоровья"],
        "delta": [
            "Блок «Быстрые действия» (3 кнопки: Записаться / Медкарта / Бонусы)",
            "Блок «Напоминания» (статика mockdata)",
            "Починить sidebar-навигацию",
        ],
    },
    {
        "title": "/cabinet/patient/appointments — Мои записи",
        "status": "404 нет | Stitch: нет | Порядок: №11",
        "has": [],
        "delta": [
            "Вкладки: Предстоящие / Прошедшие / Отменённые",
            "Карточки записей с кнопками «Отменить» / «Перенести»",
            "«Записаться снова» для прошедших",
        ],
    },
    {
        "title": "/cabinet/patient/medcard — Медицинская карта",
        "status": "Задеплоена | Stitch: ЛК пациент - Электронная медкарта | Порядок: №12",
        "has": ["Баннер, показатели, хронические, препараты, история визитов"],
        "delta": [
            "Вкладка «Документы» (загрузка PDF/JPG)",
            "QR-блок — заглушка «Скоро»",
            "Кнопка «+ Показатель» (форма: тип + значение + дата)",
            "Кнопка «Скачать PDF» — заглушка",
        ],
    },
    {
        "title": "/cabinet/patient/family — Семейный профиль",
        "status": "Задеплоена | Stitch: ЛК пациент - Семейный профиль | Порядок: №13",
        "has": ["3 карточки членов семьи, кнопка «Добавить», Family Insights"],
        "delta": [
            "Модалка добавления (ФИО + отношение + дата рождения)",
            "Максимум 5 карточек + счётчик «3/5»",
            "Рабочие ссылки «Медкарта» и «Записать»",
        ],
    },
    {
        "title": "/cabinet/patient/bonuses — Бонусная программа",
        "status": "Задеплоена | Stitch: ЛК пациена - бонусы | Порядок: №14",
        "has": ["Баланс, уровень, 4 стата, каталог наград, история транзакций"],
        "delta": [
            "Таблица 4 уровней (Bronze 5% / Silver 7% / Gold 10% / Platinum 12%)",
            "Блок «Пригласи друга» (реферальная ссылка-заглушка)",
            "Срок сгорания баллов в истории транзакций",
        ],
    },
    {
        "title": "/cabinet/patient/favorites — Избранные врачи",
        "status": "404 нет | Stitch: нет | Порядок: №15",
        "has": [],
        "delta": ["Сетка карточек (как /search) + кнопка убрать + быстрая запись"],
    },
    {
        "title": "/cabinet/patient/settings — Настройки",
        "status": "404 нет | Stitch: нет | Порядок: №16",
        "has": [],
        "delta": ["ФИО + фото, уведомления (toggle), безопасность (смена пароля), страховки"],
    },
]

for page in cabinet_patient_pages:
    add_h3(page["title"])
    add_body(page["status"])
    if page["has"]:
        add_body("Что сейчас есть:", bold=True)
        for item in page["has"]:
            add_bullet(item)
    add_body("Дельта (что делаем):", bold=True)
    for item in page["delta"]:
        add_bullet(item)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. ЛИЧНЫЙ КАБИНЕТ КЛИНИКИ
# ═══════════════════════════════════════════════════════════
add_h2("5. Личный кабинет клиники")

clinic_pages = [
    {
        "title": "/cabinet/clinic — Дашборд",
        "status": "Задеплоен | Stitch: ЛК клиника - Расширенный дашборд | Порядок: №17",
        "has": ["4 KPI, псевдографик, загрузка врачей, таблица записей"],
        "delta": [
            "Recharts — реальный линейный chart выручки",
            "Pie chart источников лидов (статика)",
            "Кнопки «Принять/Отклонить» → useState",
        ],
    },
    {
        "title": "/cabinet/clinic/doctors — Управление врачами",
        "status": "404 нет | Stitch: ЛК администартора клиники | Порядок: №18",
        "has": [],
        "delta": [
            "Таблица врачей (фото, имя, специальность, статус, рейтинг)",
            "Кнопки: Редактировать / Деактивировать",
            "Форма «Добавить врача»",
        ],
    },
    {
        "title": "/cabinet/clinic/schedule — Расписание клиники",
        "status": "404 нет | Stitch: нет | Порядок: №19",
        "has": [],
        "delta": ["Недельный вид (врачи × часы), клик на слот, экспорт"],
    },
    {
        "title": "/cabinet/clinic/appointments — Записи клиники",
        "status": "404 нет | Stitch: нет | Порядок: №20",
        "has": [],
        "delta": ["Таблица с фильтрами (дата / врач / статус), поиск по пациенту, экспорт"],
    },
    {
        "title": "/cabinet/clinic/reports — Отчёты",
        "status": "Задеплоена | Stitch: ЛК клиника - Отчет по лидам | Порядок: №21",
        "has": ["Период, 4 KPI, лиды по источникам, таблица врачей"],
        "delta": [
            "Recharts bar/line charts",
            "Retention Rate виджет",
            "Кнопка «Скачать PDF» → заглушка",
        ],
    },
    {
        "title": "/cabinet/clinic/settings — Настройки клиники",
        "status": "404 нет | Stitch: нет | Порядок: №22",
        "has": [],
        "delta": [
            "Основная информация (название, адрес, телефоны)",
            "Логотип и галерея (file input)",
            "ДМС-страховщики + Тарифный план",
        ],
    },
]

for page in clinic_pages:
    add_h3(page["title"])
    add_body(page["status"])
    if page["has"]:
        add_body("Что сейчас есть:", bold=True)
        for item in page["has"]:
            add_bullet(item)
    add_body("Дельта (что делаем):", bold=True)
    for item in page["delta"]:
        add_bullet(item)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 6. ЛИЧНЫЙ КАБИНЕТ ВРАЧА
# ═══════════════════════════════════════════════════════════
add_h2("6. Личный кабинет врача")

doctor_pages = [
    {
        "title": "/cabinet/doctor — Дашборд врача",
        "status": "Задеплоен | Stitch: ЛК ВРАЧ - Управление расписанием | Порядок: №23",
        "has": ["4 KPI, расписание на сегодня, недельный обзор, доступность, отзывы"],
        "delta": ["Recharts bar chart", "Кнопки «Принять/Отклонить» → useState", "Починить sidebar /schedule, /appointments, /patients"],
    },
    {
        "title": "/cabinet/doctor/schedule — Расписание",
        "status": "404 нет | Stitch: нет | Порядок: №24",
        "has": [],
        "delta": ["Месячный календарь с загрузкой, управление слотами, блокировка дней"],
    },
    {
        "title": "/cabinet/doctor/appointments — Приёмы",
        "status": "404 нет | Stitch: нет | Порядок: №25",
        "has": [],
        "delta": ["Вкладки Сегодня / Неделя / Архив, детали приёма, добавить назначение"],
    },
    {
        "title": "/cabinet/doctor/patients — Пациенты",
        "status": "404 нет | Stitch: нет | Порядок: №26",
        "has": [],
        "delta": ["Список + поиск, карточка пациента с историей визитов"],
    },
    {
        "title": "/cabinet/doctor/profile — Профиль (ред.)",
        "status": "404 нет | Stitch: нет | Порядок: №27",
        "has": [],
        "delta": ["Фото, ФИО, биография, образование (+добавить), сертификаты, предпросмотр"],
    },
]

for page in doctor_pages:
    add_h3(page["title"])
    add_body(page["status"])
    if page["has"]:
        add_body("Что сейчас есть:", bold=True)
        for item in page["has"]:
            add_bullet(item)
    add_body("Дельта (что делаем):", bold=True)
    for item in page["delta"]:
        add_bullet(item)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 7. БАЗА ДАННЫХ
# ═══════════════════════════════════════════════════════════
add_h2("7. База данных — основные таблицы")

db_tables = [
    ("users", "id, email, phone, role, password_hash, is_verified, created_at"),
    ("patients", "id, user_id, name, birth_date, gender, blood_type, bonus_balance, bonus_level"),
    ("doctors", "id, user_id, clinic_id, slug, specialties[], price_from, rating, accepts_dms"),
    ("clinics", "id, user_id, slug, name, address, city, lat, lng, rating, subscription_plan"),
    ("specialties", "id, name, slug, icon, parent_id"),
    ("services", "id, clinic_id, doctor_id, name, duration_minutes, price"),
    ("doctor_schedule", "id, doctor_id, day_of_week, start_time, end_time, slot_duration"),
    ("schedule_overrides", "id, doctor_id, date, type(blocked/extra), reason"),
    ("appointments", "id, patient_id, doctor_id, datetime_start, type, status, price, bonus_earned"),
    ("patient_medical_records", "id, patient_id, blood_type, allergies, chronic_conditions"),
    ("vital_signs", "id, patient_id, type, value, recorded_at"),
    ("family_members", "id, owner_patient_id, member_patient_id, relation"),
    ("reviews", "id, patient_id, doctor_id, rating, text, is_verified"),
    ("bonus_transactions", "id, patient_id, amount, type, source"),
    ("cms_articles", "id, title, slug, content, specialty_id, published_at"),
]
make_table(["Таблица", "Поля"], db_tables)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 8. BACKEND API
# ═══════════════════════════════════════════════════════════
add_h2("8. Backend API Endpoints (Фаза 2 — FastAPI)")

api_groups = {
    "Аутентификация": [
        "POST /api/v1/auth/send-sms — отправить SMS-код",
        "POST /api/v1/auth/verify-sms — подтвердить код",
        "POST /api/v1/auth/register — регистрация",
        "POST /api/v1/auth/login — вход email+пароль",
        "POST /api/v1/auth/refresh — обновить токен",
        "POST /api/v1/auth/logout — выход",
    ],
    "Врачи": [
        "GET /api/v1/doctors — список (фильтры + пагинация)",
        "GET /api/v1/doctors/{slug} — профиль",
        "GET /api/v1/doctors/{slug}/slots — доступные слоты",
        "GET /api/v1/doctors/{slug}/reviews — отзывы",
    ],
    "Клиники": [
        "GET /api/v1/clinics — список",
        "GET /api/v1/clinics/{slug} — профиль",
        "GET /api/v1/clinics/{slug}/doctors — врачи клиники",
    ],
    "Запись": [
        "POST /api/v1/appointments — создать запись",
        "GET /api/v1/appointments/{id} — детали",
        "PATCH /api/v1/appointments/{id}/cancel — отменить",
        "PATCH /api/v1/appointments/{id}/reschedule — перенести",
    ],
    "Кабинет пациента": [
        "GET /api/v1/patients/me/appointments",
        "GET /api/v1/patients/me/medical-card",
        "GET /api/v1/patients/me/bonuses",
        "GET /api/v1/patients/me/family",
    ],
    "Кабинет клиники": [
        "GET /api/v1/clinics/me/leads",
        "GET /api/v1/clinics/me/reports",
        "GET /api/v1/clinics/me/doctors",
    ],
}

for group, endpoints in api_groups.items():
    add_h4(group)
    for ep in endpoints:
        add_bullet(ep)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 9. ФУНКЦИОНАЛ СЛЕДУЮЩИХ ФАЗ
# ═══════════════════════════════════════════════════════════
add_h2("9. Функционал следующих фаз")

phase_headers = ["Категория", "Фича", "Когда"]
phase_rows = [
    ("Поиск", "Полнотекстовый поиск PostgreSQL tsvector", "Фаза 2"),
    ("Поиск", "Карта врачей (Яндекс.Карты)", "Фаза 2"),
    ("Auth", "JWT + refresh tokens", "Фаза 2"),
    ("Auth", "SMS-верификация (SMSC.ru)", "Фаза 2"),
    ("Auth", "OAuth Госуслуги", "Фаза 3"),
    ("Запись", "Реальное бронирование через API", "Фаза 2"),
    ("Запись", "SMS/email уведомления", "Фаза 2"),
    ("Запись", "Онлайн-оплата YooKassa", "Фаза 3"),
    ("Телемед", "Видеозвонок Agora.io", "Фаза 3"),
    ("SEO", "schema.org для всех профилей + sitemap.xml", "Фаза 2"),
    ("SEO", "Городские страницы /moscow/cardiologist", "Фаза 3"),
    ("Файлы", "S3 для фото и документов", "Фаза 2"),
    ("Монетизация", "Тарифы клиник + оплата", "Фаза 3"),
    ("Монетизация", "Продвижение в поиске", "Фаза 3"),
    ("Adminка", "Суперадминка /admin", "Фаза 3"),
]
make_table(phase_headers, phase_rows)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# ИТОГ
# ═══════════════════════════════════════════════════════════
add_divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Итого: 14 задеплоено ✅ | 17 создать | 4 идеи")
set_font(run, "Manrope", 12, bold=True, color=PRIMARY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Работаем постранично — каждую страницу согласовываем до начала вёрстки")
set_font(run, "Inter", 10, color=GRAY)

# Сохранение
output_path = "/Users/igor/Documents/CLAUDE CODE/Сайт медас/SITE_PLAN_v2.0.docx"
doc.save(output_path)
print(f"Сохранено: {output_path}")
