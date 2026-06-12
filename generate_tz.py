#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Стили страницы ──────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2)
section.top_margin   = Cm(2)
section.bottom_margin = Cm(2)

# ── Вспомогательные функции ─────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x0F, 0x3C, 0x78)
    return p

def para(text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Заголовок
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1A1A2E')
        shading.set(qn('w:color'), 'FFFFFF')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Строки
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg)
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
    # Ширины колонок
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ════════════════════════════════════════════════════════
#  ТИТУЛЬНАЯ СТРАНИЦА
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MEDAS')
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x3C, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Медицинский маркетплейс')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ТЕХНИЧЕСКОЕ ЗАДАНИЕ')
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Версия 1.0  •  Июнь 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('saas.med-as.ru')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x08, 0x7E, 0x8B)

doc.add_page_break()


# ════════════════════════════════════════════════════════
#  1. СУТЬ ПРОДУКТА
# ════════════════════════════════════════════════════════
heading('1. Суть продукта и цель', 1)
para('MEDAS — медицинский маркетплейс (аналог Напоправку / СберЗдоровье). Пациент находит врача или клинику, записывается онлайн, ведёт электронную медкарту и накапливает бонусы. Клиники получают управляемый поток пациентов и инструменты аналитики. Врачи управляют расписанием и видят историю своих пациентов.')

heading('Монетизация', 3)
bullet('Комиссия с каждой записи через платформу (8–15%)')
bullet('Подписка для клиник: базовый / Pro / Enterprise')
bullet('Продвижение в поиске — платное поднятие в выдаче')
bullet('B2B: корпоративное медобслуживание (ДМС-интеграция)')
hr()


# ════════════════════════════════════════════════════════
#  2. РОЛИ
# ════════════════════════════════════════════════════════
heading('2. Роли в системе', 1)
table(
    ['Роль', 'Кто это', 'Что делает'],
    [
        ['Пациент', 'Конечный пользователь', 'Ищет врача, записывается, ведёт медкарту, копит бонусы'],
        ['Врач', 'Специалист в клинике', 'Управляет расписанием, ведёт приёмы, пишет назначения'],
        ['Администратор клиники', 'Управляет клиникой', 'Добавляет врачей, видит аналитику, управляет профилем'],
        ['Контент-менеджер MEDAS', 'Сотрудник платформы', 'Наполняет каталог, проверяет клиники, пишет статьи'],
        ['Суперадмин MEDAS', 'Технический администратор', 'Полный доступ: пользователи, клиники, настройки, финансы'],
        ['Финансовый менеджер', 'Бухгалтерия / аналитика', 'Видит выручку, комиссии, выплаты клиникам'],
    ],
    [5, 5, 8]
)
hr()


# ════════════════════════════════════════════════════════
#  3. АРХИТЕКТУРА
# ════════════════════════════════════════════════════════
heading('3. Архитектура системы', 1)

heading('3.1 Стек технологий', 2)
table(
    ['Слой', 'Технология', 'Обоснование'],
    [
        ['Frontend (публичный сайт)', 'Next.js App Router + TypeScript + Tailwind CSS v4', 'SSR для SEO профилей врачей/клиник, скорость разработки'],
        ['Frontend (кабинеты / админки)', 'Next.js App Router (тот же репозиторий, отдельные роуты)', 'Единый стек, переиспользование компонентов'],
        ['Backend API', 'Python + FastAPI', 'Быстрый, async, Pydantic-валидация, OpenAPI автоматически'],
        ['Основная БД', 'PostgreSQL 16', 'ACID-транзакции, сложные JOIN, медицинские данные = строгая схема'],
        ['Кеш / очереди', 'Redis 7', 'Кеш поисковых запросов, сессии, очереди SMS/email, rate limit'],
        ['Полнотекстовый поиск', 'Elasticsearch 8 (Фаза 2)', 'Старт: pg_trgm; при 100k+ записей — миграция на ES'],
        ['Файловое хранилище', 'S3 (Selectel / Yandex)', 'Фото врачей, PDF документы, анализы'],
        ['Реверс-прокси', 'Nginx', 'SSL termination, статика, upstream к Next.js'],
        ['Деплой', 'Docker + Docker Compose + VPS', 'Простой старт, лёгкий переезд на k8s при масштабе'],
        ['CI/CD', 'GitHub Actions', 'Автосборка и деплой при push в main'],
    ],
    [4, 5, 9]
)

heading('3.2 Схема взаимодействия компонентов', 2)
code_block(
    'Браузер / Мобильное приложение\n'
    '         │\n'
    '    ┌────▼────┐\n'
    '    │  Nginx  │  ← SSL, статика, rate limit\n'
    '    └────┬────┘\n'
    '         │\n'
    '  ┌──────┴──────┐\n'
    '  │             │\n'
    '┌─▼──────┐  ┌───▼──────┐\n'
    '│Next.js │  │ FastAPI  │ ← /api/v1/*\n'
    '│:3000   │  │ :8000    │\n'
    '└────────┘  └───┬──────┘\n'
    '                │\n'
    '   ┌────────────┼──────────────┐\n'
    '   │            │              │\n'
    '┌──▼───┐  ┌─────▼──┐  ┌───────▼──────┐\n'
    '│  PG  │  │ Redis  │  │    S3 / ES   │\n'
    '└──────┘  └────────┘  └──────────────┘'
)
hr()


# ════════════════════════════════════════════════════════
#  4. БАЗА ДАННЫХ
# ════════════════════════════════════════════════════════
heading('4. Схема базы данных', 1)
para('Все таблицы создаются через Alembic-миграции. Прямые правки схемы вручную — запрещены.')

heading('4.1 Основные таблицы', 2)
db_tables = [
    ['users', 'id, email, phone, password_hash, role, is_verified, created_at, last_login, is_blocked'],
    ['patients', 'id, user_id, first_name, last_name, middle_name, birth_date, gender, blood_type, photo_url, bonus_balance, bonus_level, referral_code, referred_by'],
    ['clinics', 'id, name, slug, description, address, city, metro, lat, lng, phone, email, logo_url, working_hours(JSON), is_verified, is_premium, rating, reviews_count, subscription_plan'],
    ['doctors', 'id, user_id, clinic_id, first_name, last_name, specialties(array), experience_years, bio, photo_url, education(JSON), certificates(JSON), rating, reviews_count, price_from, is_verified, accepts_dms, is_home_visit, languages(array)'],
    ['specialties', 'id, name, slug, icon, description, parent_id'],
    ['doctor_specialties', 'doctor_id, specialty_id  ← M:M'],
    ['services', 'id, clinic_id, doctor_id, name, description, duration_minutes, price, specialty_id'],
    ['doctor_schedule', 'id, doctor_id, day_of_week(0-6), start_time, end_time, slot_duration_minutes, is_active'],
    ['schedule_overrides', 'id, doctor_id, date, start_time, end_time, type(blocked/extra), reason'],
    ['appointments', 'id, patient_id, doctor_id, clinic_id, service_id, datetime_start, datetime_end, type(clinic/online/home), status(new/confirmed/done/cancelled/no_show), price, payment_status, comment, cancel_reason, bonus_earned, created_at'],
    ['appointment_documents', 'id, appointment_id, type(diagnosis/prescription/result), content(TEXT), file_url, created_by(doctor_id)'],
    ['patient_medical_records', 'id, patient_id, blood_type, allergies(JSON), chronic_conditions(JSON), current_medications(JSON), emergency_contact(JSON)'],
    ['vital_signs', 'id, patient_id, type(pressure/weight/pulse/...), value, recorded_at, source(manual/device)'],
    ['family_members', 'id, owner_patient_id, member_patient_id, relation(spouse/child/parent)'],
    ['reviews', 'id, patient_id, doctor_id, clinic_id, appointment_id, rating(1-5), text, is_verified, created_at'],
    ['bonus_transactions', 'id, patient_id, amount, type(earned/spent/expired), source(appointment/referral/review/registration), appointment_id, description'],
    ['referrals', 'id, referrer_id, referred_id, bonus_paid, created_at'],
    ['insurance_companies', 'id, name, logo_url, is_active'],
    ['clinic_insurance', 'clinic_id, insurance_id  ← M:M'],
    ['notifications', 'id, user_id, type(sms/email/push), template, content, status, sent_at'],
    ['cms_articles', 'id, title, slug, content, author_id, specialty_id, published_at, views, is_published'],
    ['platform_settings', 'key, value, description'],
    ['audit_log', 'id, user_id, action, entity_type, entity_id, old_value(JSON), new_value(JSON), ip, created_at'],
]
table(['Таблица', 'Поля'], db_tables, [4, 14])
hr()


# ════════════════════════════════════════════════════════
#  5. ВСЕ СТРАНИЦЫ
# ════════════════════════════════════════════════════════
heading('5. Страницы и разделы системы', 1)

heading('5.1 Публичный сайт', 2)
table(
    ['URL', 'Описание'],
    [
        ['/', 'Главная — поиск, специальности, топ клиники, акции, статьи'],
        ['/search', 'Поиск врачей с фильтрами (специальность, цена, район, дата, пол, язык, ДМС)'],
        ['/search?type=clinic', 'Поиск клиник'],
        ['/doctor/[slug]', 'Профиль врача — биография, расписание, цены, отзывы, запись'],
        ['/clinic/[slug]', 'Профиль клиники — описание, врачи, галерея, отзывы, карта'],
        ['/services', 'Каталог специальностей'],
        ['/services/[slug]', 'Страница специальности (Кардиология, Неврология и т.д.)'],
        ['/articles', 'Медицинский блог'],
        ['/articles/[slug]', 'Статья'],
        ['/login', 'Вход (телефон + SMS / email + пароль)'],
        ['/register', 'Регистрация (пациент / клиника / врач)'],
        ['/forgot-password', 'Восстановление пароля'],
        ['/about', 'О платформе'],
        ['/contacts', 'Контакты'],
        ['/terms', 'Пользовательское соглашение'],
        ['/privacy', 'Политика конфиденциальности'],
    ],
    [5, 13]
)

heading('5.2 Кабинет пациента  /cabinet/patient/...', 2)
table(
    ['URL', 'Что делает'],
    [
        ['/', 'Дашборд — ближайшие записи, бонусный баланс, показатели здоровья'],
        ['/appointments', 'Все записи: предстоящие + история + отмена / перенос'],
        ['/medcard', 'Электронная медкарта: аллергии, хронические болезни, препараты'],
        ['/medcard/documents', 'Загруженные документы и анализы'],
        ['/family', 'Семейный профиль: управление медкартами до 5 членов семьи'],
        ['/bonuses', 'Бонусная программа: история начислений и трат, уровень'],
        ['/favorites', 'Избранные врачи'],
        ['/reviews', 'Мои отзывы'],
        ['/settings', 'Настройки: телефон, email, уведомления, удаление аккаунта'],
    ],
    [5, 13]
)

heading('5.3 Кабинет врача  /cabinet/doctor/...', 2)
table(
    ['URL', 'Что делает'],
    [
        ['/', 'Дашборд — расписание на сегодня, статистика приёмов'],
        ['/schedule', 'Полный календарь на месяц — управление слотами и исключениями'],
        ['/appointments', 'Все приёмы со статусами и фильтрами'],
        ['/appointments/[id]', 'Детали приёма — написать назначение, добавить диагноз, прикрепить документ'],
        ['/patients', 'База пациентов, которые были на приёме'],
        ['/patients/[id]', 'Медкарта конкретного пациента'],
        ['/profile', 'Редактирование профиля — фото, биография, образование, сертификаты'],
        ['/reviews', 'Отзывы о враче — просмотр и ответы'],
        ['/earnings', 'Выплаты и статистика доходов'],
    ],
    [5, 13]
)

heading('5.4 Кабинет клиники  /cabinet/clinic/...', 2)
table(
    ['URL', 'Что делает'],
    [
        ['/', 'Дашборд — KPI, выручка, загрузка врачей'],
        ['/doctors', 'Список врачей + добавить / редактировать / деактивировать'],
        ['/doctors/[id]', 'Управление конкретным врачом: расписание, статус, данные'],
        ['/schedule', 'Сводное расписание всех врачей клиники'],
        ['/appointments', 'Все записи клиники с фильтрами'],
        ['/patients', 'База пациентов клиники'],
        ['/reports', 'Аналитика: лиды, конверсия, выручка, бонусы, NPS'],
        ['/settings', 'Профиль клиники: контакты, фото, ДМС, описание, галерея'],
        ['/settings/subscription', 'Тарифный план и оплата'],
    ],
    [5, 13]
)

heading('5.5 Контент-админка MEDAS  /admin/content/...', 2)
para('Для сотрудников платформы, которые наполняют сайт и модерируют контент.')
table(
    ['URL', 'Что делает'],
    [
        ['/clinics', 'Список клиник — одобрить / отклонить / заблокировать / редактировать'],
        ['/clinics/[id]', 'Полное редактирование профиля клиники'],
        ['/doctors', 'Список врачей — верификация, редактирование, блокировка'],
        ['/specialties', 'Управление справочником специальностей'],
        ['/articles', 'Создание и публикация медицинских статей'],
        ['/articles/[id]', 'Редактор статьи (rich text, SEO-поля)'],
        ['/reviews', 'Модерация отзывов — одобрить / скрыть / пометить'],
        ['/seo', 'SEO-настройки ключевых страниц'],
    ],
    [5, 13]
)

heading('5.6 Суперадминка MEDAS  /admin/...', 2)
para('Полный контроль над платформой. Доступ только для технических администраторов.')
table(
    ['URL', 'Что делает'],
    [
        ['/dashboard', 'Главные метрики платформы — DAU, выручка, записи, ошибки'],
        ['/users', 'Все пользователи — поиск, блокировка, изменение роли'],
        ['/clinics', 'Все клиники с финансовыми данными и статусами подписок'],
        ['/appointments', 'Все записи на платформе с полным фильтром'],
        ['/finances', 'Комиссии, выплаты клиникам, транзакции'],
        ['/bonuses', 'Управление бонусной программой: правила, уровни, начисления'],
        ['/notifications', 'Шаблоны SMS / email уведомлений'],
        ['/settings', 'Глобальные настройки платформы: комиссии, тарифы, лимиты'],
        ['/audit', 'Лог всех действий в системе с фильтром по пользователю / сущности'],
    ],
    [5, 13]
)
hr()


# ════════════════════════════════════════════════════════
#  6. ФУНКЦИОНАЛ
# ════════════════════════════════════════════════════════
heading('6. Функционал — детально', 1)

heading('6.1 Поиск и фильтрация', 2)
bullet('Полнотекстовый поиск по имени врача, специальности, симптому, названию клиники')
bullet('Фильтры: специальность, район/метро, цена (от-до), дата доступности, пол врача, язык, ДМС, выезд на дом, онлайн-приём')
bullet('Сортировка: по рейтингу, по цене, по ближайшей доступности')
bullet('Геопоиск: «рядом со мной» (геолокация браузера)')
bullet('Старт: pg_trgm в PostgreSQL; при 100k+ записей — миграция на Elasticsearch')
bullet('Результаты кешируются в Redis на 5 минут')

heading('6.2 Система записи', 2)
bullet('Пациент выбирает врача → дату/слот → тип (клиника / онлайн / дом) → подтверждает')
bullet('Доступные слоты: doctor_schedule минус занятые appointments')
bullet('SMS-подтверждение записи — за 24 часа и за 2 часа до приёма')
bullet('Отмена: до 24 часов — бесплатно; позже — штраф (сумма задаётся в настройках)')
bullet('Повторная запись в 1 клик из истории приёмов')

heading('6.3 Авторизация и безопасность', 2)
bullet('JWT: access-токен 15 минут + refresh-токен 30 дней')
bullet('Основной способ входа: телефон + SMS-код (6 цифр)')
bullet('Альтернативный: email + пароль (bcrypt, salt 12)')
bullet('Разные JWT claims для каждой роли — патч на уровне middleware')
bullet('Rate limiting: не более 5 неудачных попыток входа с одного IP за 10 минут')
bullet('Медицинские данные доступны только владельцу и врачу на активном приёме')

heading('6.4 Электронная медкарта', 2)
bullet('Базовые данные: ФИО, группа крови, аллергии, хронические болезни, текущие препараты')
bullet('История визитов: диагноз + назначение + заметки врача (заполняет врач после приёма)')
bullet('Загрузка документов: анализы, снимки МРТ, рецепты — PDF / JPG → S3')
bullet('Показатели здоровья: давление, вес, пульс — вручную или из приложения')
bullet('Семейный профиль: управление медкартами до 5 членов семьи')
bullet('Экспорт в PDF')
bullet('QR-код для передачи врачу — временный токен на 24 часа')

heading('6.5 Бонусная программа', 2)
table(
    ['Уровень', 'Накопленные баллы', 'Кешбэк на приём'],
    [
        ['Bronze', '0 – 999', '5%'],
        ['Silver', '1 000 – 4 999', '7%'],
        ['Gold', '5 000 – 14 999', '10%'],
        ['Platinum', '15 000+', '12% + приоритетная запись'],
    ],
    [4, 5, 5]
)
bullet('Траты: скидки на приём, бесплатные анализы, приоритетная запись')
bullet('Реферальная программа: +300 баллов за приглашённого друга, который записался')
bullet('Срок жизни баллов: 365 дней с момента начисления')
bullet('Бонусы за действия: регистрация (+500), первый отзыв (+50), заполнение медкарты (+100)')

heading('6.6 Уведомления', 2)
table(
    ['Триггер', 'Канал', 'Кому'],
    [
        ['Новая запись создана', 'SMS + email', 'Пациент + Клиника'],
        ['Запись подтверждена клиникой', 'SMS', 'Пациент'],
        ['Напоминание за 24 часа', 'SMS', 'Пациент'],
        ['Напоминание за 2 часа', 'SMS', 'Пациент'],
        ['Запись отменена', 'SMS', 'Пациент + Клиника'],
        ['Запись завершена (с просьбой отзыва)', 'Push', 'Пациент'],
        ['Новая заявка на приём', 'Email', 'Администратор клиники'],
        ['Бонусы начислены', 'Push', 'Пациент'],
        ['Сертификат врача истекает (30 дней)', 'Email', 'Врач'],
    ],
    [7, 4, 7]
)

heading('6.7 Отзывы', 2)
bullet('Оставить отзыв можно только после состоявшегося приёма (appointment.status = done)')
bullet('Оценка от 1 до 5 + текст (минимум 20 символов)')
bullet('Модерация: автоматическая (фильтр нецензурной лексики) + ручная для спорных')
bullet('Врач / клиника могут ответить на отзыв')
bullet('Рейтинг пересчитывается автоматически после каждого нового отзыва')

heading('6.8 Онлайн-консультации', 2)
bullet('Видеозвонок: интеграция Agora.io (WebRTC)')
bullet('Пациент получает ссылку за 15 минут до приёма (SMS + email)')
bullet('Врач видит медкарту пациента прямо во время звонка (отдельная панель)')
bullet('После окончания — врач заполняет назначение и закрывает приём')

heading('6.9 Аналитика для клиники', 2)
bullet('Источники лидов: из поиска, из каталога, прямая ссылка, реклама')
bullet('Конверсия: просмотр профиля → запись → состоявшийся приём')
bullet('Выручка по дням / неделям / месяцам с графиком')
bullet('Загрузка врачей: % занятости по слотам')
bullet('Возвращаемость пациентов (Retention Rate)')
bullet('NPS и средний рейтинг по отзывам')
hr()


# ════════════════════════════════════════════════════════
#  7. ИНТЕГРАЦИИ
# ════════════════════════════════════════════════════════
heading('7. Внешние интеграции', 1)
table(
    ['Сервис', 'Для чего', 'Альтернатива'],
    [
        ['СМSC / SMS.ru', 'Отправка SMS-кодов и уведомлений', 'МТС Коммуникатор'],
        ['SendGrid / Unisender', 'Email-рассылки транзакционные и маркетинговые', 'Mailgun'],
        ['Яндекс.Карты API', 'Карта на страницах врача / клиники, геопоиск', '2GIS API'],
        ['Agora.io', 'Видеозвонки (WebRTC) для онлайн-консультаций', 'VONAGE'],
        ['ЮKassa / Tinkoff Pay', 'Онлайн-оплата приёмов, подписок', 'Robokassa'],
        ['Selectel Object Storage', 'Хранение фото, PDF документов, анализов', 'Yandex Object Storage'],
        ['Firebase / OneSignal', 'Push-уведомления в браузере и мобильном', '–'],
        ['Sentry', 'Мониторинг ошибок в реальном времени', 'GlitchTip (self-hosted)'],
        ['Prometheus + Grafana', 'Метрики производительности и алерты', '–'],
    ],
    [5, 7, 6]
)
hr()


# ════════════════════════════════════════════════════════
#  8. БЕЗОПАСНОСТЬ
# ════════════════════════════════════════════════════════
heading('8. Безопасность и соответствие законодательству', 1)

heading('8.1 Законодательство', 2)
bullet('ФЗ-152 «О персональных данных» — шифрование, согласие на обработку, возможность удаления данных по запросу')
bullet('ФЗ-323 «Об основах охраны здоровья» — медицинские данные хранятся только в РФ (хостинг на Selectel / Timeweb)')
bullet('Обязательный аудит-лог всех действий с медицинскими данными (таблица audit_log)')
bullet('Форма согласия на обработку персональных данных при регистрации (юридически оформленная)')

heading('8.2 Технические меры', 2)
bullet('HTTPS везде — Let\'s Encrypt, авторенью через cron')
bullet('Медицинские данные шифруются at-rest (PostgreSQL column-level encryption или pgcrypto)')
bullet('Пароли: bcrypt с cost-factor 12')
bullet('SQL-инъекции: только ORM (SQLAlchemy), никакого raw SQL с пользовательским вводом')
bullet('XSS: sanitize всего пользовательского контента перед сохранением в БД')
bullet('Регулярный бэкап БД: ежедневно → Selectel S3, retention 30 дней')
bullet('Раздельный доступ: врач видит только своих пациентов, клиника — только свою')
bullet('Rate limiting на все API-эндпоинты через Redis')
hr()


# ════════════════════════════════════════════════════════
#  9. МОБИЛЬНОЕ ПРИЛОЖЕНИЕ
# ════════════════════════════════════════════════════════
heading('9. Мобильное приложение (Фаза 3)', 1)
para('На старте — адаптивный веб-сайт (мобильная версия). После MVP — нативные приложения:')
bullet('React Native — единый код для iOS + Android')
bullet('Push-уведомления о записях, результатах анализов и бонусах')
bullet('Биометрический вход (Face ID / Touch ID)')
bullet('Виджет на экране блокировки с ближайшей записью')
bullet('Интеграция с Apple Health / Google Health — автоимпорт показателей здоровья')
bullet('Оффлайн-режим для медкарты (кеш последних данных)')
hr()


# ════════════════════════════════════════════════════════
#  10. ПЛАН РАЗРАБОТКИ (ФАЗЫ)
# ════════════════════════════════════════════════════════
heading('10. Очерёдность разработки', 1)

heading('Фаза 1 — MVP (2–3 месяца)', 2)
table(
    ['№', 'Задача', 'Приоритет'],
    [
        ['1', 'Авторизация: SMS + email для 3 ролей (пациент / клиника / врач)', 'P0'],
        ['2', 'Профили клиник и врачей + контент-админка для наполнения', 'P0'],
        ['3', 'Поиск + фильтры (pg_trgm)', 'P0'],
        ['4', 'Запись на приём (без онлайн-оплаты, подтверждение вручную клиникой)', 'P0'],
        ['5', 'Кабинет пациента: записи + базовая медкарта', 'P0'],
        ['6', 'Кабинет клиники: входящие заявки, управление расписанием', 'P0'],
        ['7', 'Отзывы с модерацией', 'P1'],
        ['8', 'SMS-уведомления (новая запись, напоминание за 24 ч)', 'P1'],
    ],
    [1, 12, 5]
)

heading('Фаза 2 — Монетизация (1–2 месяца)', 2)
table(
    ['№', 'Задача', 'Приоритет'],
    [
        ['1', 'Онлайн-оплата (ЮKassa)', 'P0'],
        ['2', 'Бонусная программа (4 уровня, реферальная)', 'P0'],
        ['3', 'Подписки для клиник (базовый / Pro / Enterprise)', 'P0'],
        ['4', 'Продвижение клиник в поиске (платное поднятие)', 'P1'],
        ['5', 'Аналитика для клиники (полный дашборд)', 'P1'],
    ],
    [1, 12, 5]
)

heading('Фаза 3 — Рост (2 месяца)', 2)
table(
    ['№', 'Задача', 'Приоритет'],
    [
        ['1', 'Онлайн-консультации (Agora.io WebRTC)', 'P0'],
        ['2', 'Семейный профиль и медкарта для членов семьи', 'P1'],
        ['3', 'Загрузка документов в медкарту (S3)', 'P1'],
        ['4', 'React Native мобильное приложение', 'P1'],
        ['5', 'Интеграция с ДМС', 'P2'],
        ['6', 'Корпоративные клиенты (B2B)', 'P2'],
    ],
    [1, 12, 5]
)

heading('Фаза 4 — Масштабирование', 2)
bullet('Миграция поиска на Elasticsearch')
bullet('Горизонтальное масштабирование API (несколько инстансов FastAPI за балансировщиком)')
bullet('CDN для статики и фотографий врачей / клиник')
bullet('Kubernetes вместо Docker Compose при нагрузке > 10k DAU')
bullet('Агрегация и импорт клиник / врачей из внешних источников (2GIS, Яндекс.Карты)')
hr()


# ════════════════════════════════════════════════════════
#  11. РИСКИ
# ════════════════════════════════════════════════════════
heading('11. Ключевые риски', 1)
table(
    ['Риск', 'Вероятность', 'Митигация'],
    [
        ['Клиники не хотят регистрироваться (проблема курицы и яйца)',
         'Высокая', 'Первые 30 клиник — вручную с менеджером, 3 месяца бесплатно'],
        ['Мало реальных отзывов на старте',
         'Средняя', 'Приоритет врачам с отзывами; email-кампания для пациентов после приёма'],
        ['Нарушение ФЗ-152 (персональные данные)',
         'Средняя', 'Юрист для разработки политики; хостинг в РФ с первого дня'],
        ['Высокий no-show rate (пациент не приходит)',
         'Высокая', 'Штраф за позднюю отмену + SMS-напоминания за 24 ч и 2 ч'],
        ['Конкуренция с СберЗдоровье / Напоправку',
         'Средняя', 'Нишевание: конкретный регион, лучший UX, бонусная программа'],
        ['Взлом и утечка медицинских данных',
         'Низкая', 'Шифрование at-rest, аудит-лог, penetration testing перед запуском'],
    ],
    [6, 3, 9]
)
hr()


# ════════════════════════════════════════════════════════
#  12. КОМАНДА
# ════════════════════════════════════════════════════════
heading('12. Команда для реализации', 1)
table(
    ['Роль', 'Кол-во', 'Задачи'],
    [
        ['Fullstack-разработчик', '1–2', 'Next.js (frontend) + FastAPI (backend), основная разработка'],
        ['DevOps', '0.5 (part-time)', 'VPS, CI/CD, мониторинг, SSL, бэкапы'],
        ['UX/UI-дизайнер', '1', 'Макеты страниц, дизайн-система, мобильная адаптация'],
        ['Контент-менеджер', '1', 'Наполнение базы клиник и врачей, медицинские статьи'],
        ['QA-тестировщик', '0.5 (part-time)', 'Тестирование перед каждым релизом'],
    ],
    [5, 3, 10]
)
para('Весь стек (Next.js + FastAPI) позволяет вести разработку силами 1–2 разработчиков с поддержкой Claude Code.', bold=False)
hr()


# ════════════════════════════════════════════════════════
#  13. ТЕКУЩЕЕ СОСТОЯНИЕ
# ════════════════════════════════════════════════════════
heading('13. Текущее состояние проекта', 1)
table(
    ['Компонент', 'Статус', 'Примечание'],
    [
        ['VPS (85.239.44.14)', '✅ Готово', 'Timeweb, root-доступ по SSH-ключу'],
        ['SSL-сертификат', '✅ Готово', 'Let\'s Encrypt, действует до 2026-09-09, авторенью'],
        ['Docker + Nginx', '✅ Готово', 'docker-compose.yml, HTTPS → frontend:3000'],
        ['Frontend (Next.js)', '✅ Запущен', '14 статичных страниц, данные захардкожены'],
        ['GitHub Actions CI/CD', '⚠️ Настроен', 'Нужно добавить секреты VPS_HOST + VPS_SSH_KEY'],
        ['Backend API', '❌ Не создан', 'Нужно создать с Фазы 1'],
        ['База данных PostgreSQL', '❌ Не создана', 'Нужно развернуть на VPS'],
        ['Redis', '❌ Не создан', 'Нужно добавить в docker-compose'],
        ['Авторизация', '❌ Не реализована', 'Фаза 1, задача 1'],
        ['Кабинет пациента', '❌ Не реализован', 'Фаза 1, задача 5'],
        ['Кабинет клиники', '❌ Не реализован', 'Фаза 1, задача 6'],
        ['Контент-админка', '❌ Не реализована', 'Фаза 1, задача 2'],
    ],
    [5, 3, 10]
)


# ════════════════════════════════════════════════════════
#  ФИНАЛЬНАЯ СТРАНИЦА
# ════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\nMEDAS — Техническое задание v1.0')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Июнь 2026  •  saas.med-as.ru')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


# ── Сохранить ────────────────────────────────────────────
output_path = '/Users/igor/Documents/CLAUDE CODE/Сайт медас/MEDAS_ТЗ_v1.0.docx'
doc.save(output_path)
print(f"✅ Документ сохранён: {output_path}")
