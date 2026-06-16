#!/usr/bin/env python3
"""Export MASTER_TZ.md v2.0 to MASTER_TZ.docx — полный аудит платформы."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

DOCX_PATH = "/Users/igor/Documents/CLAUDE CODE/Сайт медас/MASTER_TZ.docx"

NAVY = RGBColor(0x00, 0x30, 0x87)
GREEN = RGBColor(0x00, 0xa9, 0x82)
GRAY = RGBColor(0x60, 0x60, 0x60)
RED = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xFF, 0x66, 0x00)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = GREEN
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    elif level == 4:
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
        run.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def add_para(text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '003087')
        tcPr.append(shd)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════
# ТИТУЛЬНАЯ СТРАНИЦА
# ══════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
r = title_p.add_run("MEDAS")
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("Техническое задание — Полный аудит платформы")
r2.font.size = Pt(18)
r2.font.color.rgb = GRAY

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = ver_p.add_run(f"Версия 2.0  |  {datetime.date.today().strftime('%d.%m.%Y')}  |  Конфиденциально")
r3.font.size = Pt(11)
r3.font.color.rgb = GRAY

doc.add_paragraph()
note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = note_p.add_run("Документ основан на реальном аудите кода — каждое утверждение проверено по файлам проекта")
r4.font.size = Pt(10)
r4.font.italic = True
r4.font.color.rgb = GREEN

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. ОБЩИЙ СТАТУС
# ══════════════════════════════════════════════════════════════

add_heading("1. Общий статус платформы", level=1)

add_heading("Инфраструктура", level=2)
add_table(
    ["Параметр", "Значение"],
    [
        ("Frontend", "saas.med-as.ru (Next.js 15 App Router)"),
        ("API", "api.med-as.ru (FastAPI Python 3.12)"),
        ("VPS", "85.239.44.14"),
        ("БД", "PostgreSQL 16 + Redis + MinIO"),
        ("Деплой", "GitHub Actions → docker build + docker cp"),
        ("SSL", "saas.med-as.ru истекает 2026-09-11, certbot crontab ✅"),
    ]
)

add_heading("Данные в базе (на 2026-06-16)", level=2)
add_table(
    ["Сущность", "Количество", "Примечание"],
    [
        ("Врачи", "6", "Только 3 имеют страницы профиля на сайте"),
        ("Клиники", "5", "Только 2 имеют страницы профиля на сайте"),
        ("Пользователи", "N", "Пациенты + 1 тест-клиника + тест-врачи"),
        ("Записи (appointments)", "433+", "Seed-данные для тестирования дашборда"),
        ("Бонусные транзакции", "N", "Welcome +500 при регистрации"),
    ]
)

add_heading("Итоговые цифры аудита", level=2)
add_table(
    ["Метрика", "Значение"],
    [
        ("Всего страниц", "22"),
        ("Страниц полностью на реальных данных", "8"),
        ("Страниц частично на реальных данных", "5"),
        ("Страниц только на mock/статичных данных", "5"),
        ("Страниц-заглушек (404)", "4"),
        ("Критических broken flows (сломанных сценариев)", "16"),
        ("Декоративных кнопок (ничего не делают)", "11"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. МАТРИЦА РЕАЛЬНЫЕ VS ЗАГЛУШКИ
# ══════════════════════════════════════════════════════════════

add_heading("2. Матрица: Реальные данные vs Заглушки", level=1)

add_heading("По страницам", level=2)
add_table(
    ["URL", "Реальный API", "Заглушка/Статика", "Сломано"],
    [
        ("/", "—", "Всё (stats, doctors, clinics, articles)", "Clinic cards → 404"),
        ("/search", "GET /doctors (6 врачей)", "Фильтры metro/DMS/online", "3/6 doctor links → 404"),
        ("/doctors", "—", "Всё (SEO-статика)", "—"),
        ("/doctor/[slug]", "slots, available-days", "Профиль из lib/doctors.ts", "3/6 API-врачей → 404"),
        ("/doctor/[slug]/booking", "POST /appointments, slots", "—", "—"),
        ("/clinics", "GET /clinics (5 клиник)", "—", "3/5 clinic links → 404"),
        ("/clinic/[slug]", "—", "Весь профиль (lib/clinics.ts)", "3/5 API-слагов → 404"),
        ("/login", "POST /auth/*", "—", "Реальные звонки pending"),
        ("/register", "POST /auth/*", "—", "Реальные звонки pending"),
        ("/services", "GET /doctors", "—", "3/6 doctor links → 404"),
        ("/about", "—", "Всё (статика)", "—"),
        ("/cabinet/patient", "GET /auth/me, /appointments", "Бонусы, health stats", "Nav → 404"),
        ("/cabinet/patient/bonuses", "GET /auth/me, /bonuses/my", "Rewards секция", '"Получить" → ничего'),
        ("/cabinet/patient/medcard", "—", "Всё (чужие данные!)", "PDF, Share → ничего"),
        ("/cabinet/patient/family", "—", "Всё (чужие данные!)", "Все кнопки → ничего"),
        ("/cabinet/patient/appointments", "—", "—", "❌ 404 СТРАНИЦА НЕ СОЗДАНА"),
        ("/cabinet/patient/favorites", "—", "—", "❌ 404 СТРАНИЦА НЕ СОЗДАНА"),
        ("/cabinet/clinic", "GET /appointments/clinic/stats", "—", "—"),
        ("/cabinet/clinic/appointments", "GET /appointments/clinic", "—", "—"),
        ("/cabinet/clinic/doctors", "GET /clinics/{id}/doctors", "—", "—"),
        ("/cabinet/clinic/schedule", "GET+POST+DELETE /day-offs", "—", "—"),
        ("/cabinet/clinic/reports", "GET /analytics", "—", "—"),
        ("/cabinet/clinic/settings", "GET /auth/me (phone)", "userName hardcoded", "Нельзя изменить данные"),
        ("/cabinet/doctor", "GET /appointments/doctor, /auth/me", "—", "Nav → 404"),
        ("/cabinet/doctor/schedule", "—", "—", "❌ 404 СТРАНИЦА НЕ СОЗДАНА"),
        ("/cabinet/doctor/settings", "—", "—", "❌ 404 СТРАНИЦА НЕ СОЗДАНА"),
    ]
)

add_heading("Врачи: реальные слаги vs профили", level=2)
add_table(
    ["Слаг (API)", "Имя", "Специальность", "/doctor/[slug]", "/search список"],
    [
        ("anna-sokolova", "Соколова Анна Михайловна", "Кардиолог", "✅ Работает", "✅"),
        ("igor-petrov", "Петров Игорь Сергеевич", "Хирург", "✅ Работает", "✅"),
        ("maria-kozlova", "Козлова Мария Александровна", "Педиатр", "✅ Работает", "✅"),
        ("elena-morozova-sm", "Морозова Елена Владимировна", "Невролог", "❌ 404", "✅"),
        ("pavel-ivanov-sm", "Иванов Павел Сергеевич", "Терапевт", "❌ 404", "✅"),
        ("aleksey-sidorov-sm", "Сидоров Алексей Николаевич", "Гастроэнтеролог", "❌ 404", "✅"),
    ]
)
add_para(
    "Причина: /doctor/[slug]/page.tsx читает из статического файла lib/doctors.ts (3 записи). "
    "Три API-врача там отсутствуют → notFound() → 404.\n"
    "Исправление: Заменить getDoctorBySlug(slug) на fetchDoctorBySlug(slug) из lib/api.ts.",
    italic=True
)

add_heading("Клиники: реальные слаги vs профили", level=2)
add_table(
    ["Слаг (API)", "Название", "/clinic/[slug]", "/clinics список"],
    [
        ("stomatologiya-ulybka", "Стоматология Улыбка", "✅ Работает", "✅"),
        ("semeynyy-doktor", "Семейный доктор", "✅ Работает", "✅"),
        ("medicina-na-tsvetnoy", "Медицина на Цветном", "❌ 404", "✅"),
        ("evromedservice", "ЕвроМедСервис", "❌ 404", "✅"),
        ("sm-klinika", "СМ-Клиника", "❌ 404", "✅"),
    ]
)

add_heading("Фильтры поиска — что работает", level=2)
add_table(
    ["Фильтр", "Для реальных врачей", "Причина"],
    [
        ("Поиск по имени/специальности (q)", "✅ Работает", "Текстовый match по API-данным"),
        ("Фильтр по цене", "✅ Работает", "price из API"),
        ("Фильтр по рейтингу", "✅ Работает", "rating из API"),
        ("Фильтр по метро", "❌ 0 результатов", "apiDoctorToDoctor маппит metro: []"),
        ("Фильтр ДМС", "❌ 0 результатов", "apiDoctorToDoctor: acceptsDMS: false"),
        ("Фильтр онлайн-приём", "❌ 0 результатов", "apiDoctorToDoctor: online: false"),
        ("Фильтр выезд на дом", "❌ 0 результатов", "apiDoctorToDoctor: homeVisit: false"),
        ("Фильтр по полу", "❌ Не работает", "gender не маппируется"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. АУДИТ ВСЕХ СТРАНИЦ
# ══════════════════════════════════════════════════════════════

add_heading("3. Полный аудит страниц", level=1)

# Grouped summary
add_heading("Публичные страницы", level=2)
add_table(
    ["URL", "Данные", "Статус", "Ключевые проблемы"],
    [
        ("/", "Mock", "⚠️ Частично", "Клиники на главной → 404"),
        ("/search", "API врачи", "⚠️ Частично", "Фильтры ДМС/метро/онлайн сломаны, 3/6 врачей → 404"),
        ("/doctors", "Статика", "✅ ОК", "Исправлено: Header/Footer, ссылки специальностей"),
        ("/clinics", "API клиники", "⚠️ Частично", "3/5 клиник при клике → 404"),
        ("/doctor/[slug]", "Смешанные", "⚠️ Частично", "3/6 API-врачей → 404"),
        ("/doctor/[slug]/booking", "API", "✅ ОК", "Исправлено: pre-fill даты/времени"),
        ("/clinic/[slug]", "Статика", "⚠️ Частично", "3/5 API-слагов → 404"),
        ("/login", "API", "✅ ОК", "Звонки pending (мастер-код работает)"),
        ("/register", "API", "✅ ОК", "Аналогично"),
        ("/services", "API врачи", "⚠️ Частично", "3/6 врачей → 404"),
        ("/about", "Статика", "✅ ОК", "Полностью статичная"),
    ]
)

add_heading("ЛК Пациента", level=2)
add_table(
    ["URL", "Реальные данные", "Что сломано / Mock"],
    [
        ("/cabinet/patient", "Имя (auth/me), Записи (appointments)", "Бонусы 1230 hardcoded, Health stats hardcoded"),
        ("/cabinet/patient/bonuses", "Баланс, история транзакций", "Кнопки 'Получить' → ничего"),
        ("/cabinet/patient/medcard", "НИЧЕГО", "Всё hardcoded: 'Алекс Стерлинг', 3 чужих визита"),
        ("/cabinet/patient/family", "НЕЧЕГО", "Всё hardcoded: 'Семья Стерлинг' из 3 членов"),
        ("/cabinet/patient/appointments", "—", "❌ СТРАНИЦА ОТСУТСТВУЕТ → 404"),
        ("/cabinet/patient/favorites", "—", "❌ СТРАНИЦА ОТСУТСТВУЕТ → 404"),
    ]
)

add_heading("ЛК Клиники", level=2)
add_table(
    ["URL", "Данные", "Статус"],
    [
        ("/cabinet/clinic", "KPI + воронка + timeline + выручка", "✅ Всё реальное"),
        ("/cabinet/clinic/appointments", "Таблица записей + CSV", "✅ Всё реальное"),
        ("/cabinet/clinic/doctors", "Список врачей + цены", "✅ Всё реальное"),
        ("/cabinet/clinic/schedule", "Расписание + DoctorDayOff", "✅ Всё реальное"),
        ("/cabinet/clinic/reports", "Аналитика по типам/врачам/бонусам", "✅ Всё реальное"),
        ("/cabinet/clinic/settings", "Телефон из API, userName hardcoded", "⚠️ Нет редактирования"),
    ]
)

add_heading("ЛК Врача", level=2)
add_table(
    ["URL", "Данные", "Статус"],
    [
        ("/cabinet/doctor", "KPI + DayTimeline + таблица записей", "✅ Всё реальное"),
        ("/cabinet/doctor/schedule", "—", "❌ СТРАНИЦА ОТСУТСТВУЕТ → 404"),
        ("/cabinet/doctor/settings", "—", "❌ СТРАНИЦА ОТСУТСТВУЕТ → 404"),
    ]
)

add_heading("Декоративные кнопки (ничего не делают)", level=2)
add_table(
    ["Страница", "Кнопка", "Что должна делать"],
    [
        ("/cabinet/patient/medcard", "Скачать PDF", "Генерировать PDF медкарты"),
        ("/cabinet/patient/medcard", "Поделиться", "Отправить ссылку/PDF"),
        ("/cabinet/patient/medcard", "Документы (на визите)", "Показать документы визита"),
        ("/cabinet/patient/medcard", "Фильтр специальности", "Фильтровать историю"),
        ("/cabinet/patient/bonuses", "Получить (x4 кнопки)", "Активировать награду"),
        ("/cabinet/patient/family", "Добавить члена семьи", "Форма добавления"),
        ("/cabinet/patient/family", "Медкарта (x3 члена)", "Открыть медкарту члена"),
        ("/cabinet/patient/family", "Записать (x3 члена)", "Записать члена к врачу"),
        ("/doctor/[slug]", "В избранное (♡)", "Добавить в /favorites"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. BROKEN FLOWS
# ══════════════════════════════════════════════════════════════

add_heading("4. Критические Broken Flows (сломанные сценарии)", level=1)
add_para(
    "Broken flow — это сценарий, который пользователь ожидает выполнить, но получает 404 или видит неверные данные. "
    "Все 16 пунктов ниже проверены реальным запросом к сайту или чтением кода.",
    italic=True
)

add_table(
    ["#", "Путь", "Причина (в коде)", "Затронуто", "Приоритет"],
    [
        ("BF-01", "/ → клиника → 404", "ClinicsSection: фиктивные IDs, маршрут /clinics/[slug] не существует", "100% гостей главной", "P0"),
        ("BF-02", "/search → врач → 404", "lib/doctors.ts: только 3 из 6 API-врачей, getClinicBySlug → notFound()", "50% врачей (3/6)", "P0"),
        ("BF-03", "/clinics → клиника → 404", "lib/clinics.ts: 3 из 5 API-клиник не зарегистрированы", "60% клиник (3/5)", "P0"),
        ("BF-04", "ЛК пациента → Приёмы → 404", "Нет файла page.tsx по этому маршруту", "100% пациентов", "P0"),
        ("BF-05", "ЛК пациента → Избранные → 404", "Нет страницы и нет модели в БД", "100% пациентов", "P1"),
        ("BF-06", "ЛК врача → Расписание → 404", "Нет файла page.tsx", "100% врачей", "P0"),
        ("BF-07", "ЛК врача → Настройки → 404", "Нет файла page.tsx", "100% врачей", "P1"),
        ("BF-08", '/search → фильтр "ДМС" → 0 результатов', "apiDoctorToDoctor: acceptsDMS всегда false", "100% пользователей фильтров", "P1"),
        ("BF-09", "/search → фильтр метро → 0 результатов", "apiDoctorToDoctor: metro всегда []", "100% пользователей фильтров", "P1"),
        ("BF-10", "Медкарта → чужие данные", "Весь контент hardcoded (Алекс Стерлинг)", "100% пациентов", "P0"),
        ("BF-11", "Семья → чужие данные", "Весь контент hardcoded (Семья Стерлинг)", "100% пациентов", "P1"),
        ("BF-12", "Бонусы на главной ЛК → неверная сумма", "Число 1230 hardcoded, нет GET /auth/me для виджета", "100% пациентов", "P1"),
        ("BF-13", 'Бонусы "Получить" → ничего', "onClick отсутствует, нет модели BonusReward", "100% пациентов", "P2"),
        ("BF-14", '"Скачать PDF" медкарты → ничего', "onClick отсутствует", "100% пациентов", "P2"),
        ("BF-15", '"В избранное" на враче → ничего', "onClick отсутствует, нет Favorite-модели", "100% пользователей", "P1"),
        ("BF-16", "Настройки клиники → нельзя изменить", "Нет формы, нет PATCH /clinics/{id}", "100% владельцев клиник", "P1"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. СЦЕНАРИИ ПОЛЬЗОВАТЕЛЕЙ
# ══════════════════════════════════════════════════════════════

add_heading("5. Детальные пользовательские сценарии", level=1)

# ────────── Пациент ──────────
add_heading("5.1 Пациент — поиск и запись к врачу (полный путь)", level=2)
add_para(
    "Что должно происходить по задумке, и что происходит в реальности — шаг за шагом.",
    italic=True
)

add_table(
    ["Шаг", "Страница", "API вызов", "Что видит пациент", "Статус"],
    [
        ("1", "/ (главная)", "—", "Поиск, секции врачей и клиник", "⚠️ Клиники → 404"),
        ("2", "/search?q=кардиолог", "GET /doctors (SSR)", "6 карточек, фильтр по q работает", "⚠️ Фильтры ДМС/метро сломаны"),
        ("3", "/doctor/anna-sokolova", "GET /doctors/anna-sokolova/available-days\nGET /doctors/anna-sokolova/slots", "Профиль врача + реальный календарь", "✅"),
        ("3b", "/doctor/elena-morozova-sm", "—", "Страница не найдена", "❌ 404 (BF-02)"),
        ("4", "/doctor/anna-sokolova/booking?date=...&time=...", "—", "Форма с предзаполненной датой и временем", "✅ (исправлено)"),
        ("5", "Отправка формы", "POST /appointments", "Запись создана, redirect → ЛК", "✅"),
        ("6", "/cabinet/patient", "GET /auth/me\nGET /appointments/patient", "Реальные записи, имя из API", "⚠️ Бонусы hardcoded"),
        ("7", "/cabinet/patient/appointments", "—", "Страница не найдена", "❌ 404 (BF-04)"),
        ("8", "/cabinet/patient/bonuses", "GET /auth/me\nGET /bonuses/my", "Реальный баланс, история транзакций", "✅"),
        ("9", "/cabinet/patient/medcard", "—", "Чужие данные (Алекс Стерлинг)", "❌ Mock"),
        ("10", "/cabinet/patient/favorites", "—", "Страница не найдена", "❌ 404 (BF-05)"),
    ]
)

# ────────── Клиника ──────────
add_heading("5.2 Владелец клиники — управление (полный путь)", level=2)
add_para("Тест-аккаунт: +70000000001 / 123456", italic=True)

add_table(
    ["Шаг", "Страница", "API вызов", "Что видит", "Статус"],
    [
        ("1", "/login", "POST /auth/login\nPOST /auth/verify-otp\nGET /auth/me", "OTP форма → JWT → role=clinic", "✅"),
        ("2", "/cabinet/clinic", "GET /appointments/clinic/stats", "KPI + воронка + timeline + выручка", "✅ Полностью реально"),
        ("3", "/cabinet/clinic/appointments", "GET /appointments/clinic", "Таблица всех записей + фильтры", "✅"),
        ("4", "/cabinet/clinic/doctors", "GET /clinics/{id}/doctors", "Список врачей, inline edit цены", "✅"),
        ("5", "/cabinet/clinic/schedule", "GET+POST+DELETE /doctors/{id}/day-offs", "Расписание + блокировка дат", "✅"),
        ("6", "/cabinet/clinic/reports", "GET /appointments/clinic/analytics", "Аналитика: типы, врачи, бонусы", "✅"),
        ("7", "/cabinet/clinic/settings", "GET /auth/me", "Read-only, userName hardcoded", "⚠️ Нет редактирования"),
    ]
)

# ────────── Врач ──────────
add_heading("5.3 Врач — работа с приёмами", level=2)

add_table(
    ["Шаг", "Страница", "API вызов", "Что видит", "Статус"],
    [
        ("1", "/login", "POST /auth/*, GET /auth/me", "JWT, role=doctor", "✅"),
        ("2", "/cabinet/doctor", "GET /appointments/doctor\nGET /auth/me", "KPI, DayTimeline, таблица записей", "✅ Всё реально"),
        ("3", "/cabinet/doctor/schedule", "—", "Страница не найдена", "❌ 404 (BF-06)"),
        ("4", "/cabinet/doctor/settings", "—", "Страница не найдена", "❌ 404 (BF-07)"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. КАРТА API
# ══════════════════════════════════════════════════════════════

add_heading("6. Карта API — все endpoints", level=1)

add_heading("Реализованные (проверено по коду)", level=2)
add_table(
    ["Endpoint", "Метод", "Используется где", "Статус"],
    [
        ("/auth/register", "POST", "/register (RegisterForm)", "✅"),
        ("/auth/login", "POST", "/login (LoginForm)", "✅"),
        ("/auth/verify-otp", "POST", "/login, /register", "✅"),
        ("/auth/me", "GET", "PatientHeroGreeting, bonuses, doctor, clinic/settings", "✅"),
        ("/doctors", "GET", "/search (SSR), /services", "✅"),
        ("/doctors/{slug}", "GET", "НЕ используется frontend — читает lib/doctors.ts", "✅ API / ❌ frontend"),
        ("/doctors/{slug}/slots", "GET", "AppointmentSidebarV2, BookingForm", "✅"),
        ("/doctors/{slug}/available-days", "GET", "AppointmentCalendar", "✅"),
        ("/doctors/{id}/schedule", "GET", "/cabinet/clinic/schedule", "✅"),
        ("/doctors/{id}/day-offs", "POST", "/cabinet/clinic/schedule", "✅"),
        ("/doctors/{id}/day-offs/{date}", "DELETE", "/cabinet/clinic/schedule", "✅"),
        ("/clinics", "GET", "/clinics (SSR)", "✅"),
        ("/clinics/{slug}", "GET", "НЕ используется frontend — читает lib/clinics.ts", "✅ API / ❌ frontend"),
        ("/clinics/{clinic_id}/doctors", "GET", "/cabinet/clinic/doctors", "✅"),
        ("/bonuses/my", "GET", "/cabinet/patient/bonuses", "✅"),
        ("/appointments", "POST", "BookingForm", "✅"),
        ("/appointments/patient", "GET", "PatientAppointments", "✅"),
        ("/appointments/clinic", "GET", "/cabinet/clinic/appointments", "✅"),
        ("/appointments/clinic/stats", "GET", "/cabinet/clinic dashboard", "✅"),
        ("/appointments/clinic/analytics", "GET", "/cabinet/clinic/reports", "✅"),
        ("/appointments/doctor", "GET", "/cabinet/doctor", "✅"),
        ("/appointments/{id}/confirm", "PATCH", "ClinicAppointments, sidebar", "✅"),
        ("/appointments/{id}/cancel", "PATCH", "PatientAppointments, ClinicAppointments", "✅"),
    ]
)

add_heading("Отсутствующие (нужно добавить)", level=2)
add_table(
    ["Endpoint", "Метод", "Нужен для", "Приоритет"],
    [
        ("/favorites/my", "GET", "/cabinet/patient/favorites", "P1"),
        ("/favorites/{doctor_slug}", "POST", "Кнопка ♡ на /doctor/[slug]", "P1"),
        ("/favorites/{doctor_slug}", "DELETE", "Убрать из избранного", "P1"),
        ("/clinics/{id}", "PATCH", "/cabinet/clinic/settings редактирование", "P1"),
        ("/doctors/{id}/profile", "PATCH", "/cabinet/doctor/settings", "P1"),
        ("/doctors/{id}/profile", "GET", "/cabinet/doctor/settings", "P1"),
        ("/appointments/bonus-redeem", "POST", "Кнопки Получить в /cabinet/patient/bonuses", "P2"),
        ("/medrecords/my", "GET", "/cabinet/patient/medcard реальные данные", "P2"),
        ("/family-members", "GET/POST", "/cabinet/patient/family реальные данные", "P2"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. ПЛАН ИСПРАВЛЕНИЙ
# ══════════════════════════════════════════════════════════════

add_heading("7. Приоритизированный план исправлений", level=1)

add_heading("P0 — Критические (блокируют пользователей прямо сейчас)", level=2)
add_para(
    "Без этих исправлений платформу нельзя показывать клиентам или инвесторам. "
    "Любой стандартный путь пользователя заканчивается 404.",
    italic=True
)

add_table(
    ["#", "Задача", "Файл (изменить)", "Backend готов?", "Сложность"],
    [
        ("P0-1", "Профили 3 врачей: elena-morozova-sm, pavel-ivanov-sm, aleksey-sidorov-sm",
         "frontend/src/app/doctor/[slug]/page.tsx\n→ getDoctorBySlug → fetchDoctorBySlug", "✅ GET /doctors/{slug}", "M"),
        ("P0-2", "Профили 3 клиник: medicina-na-tsvetnoy, evromedservice, sm-klinika",
         "frontend/src/app/clinic/[slug]/page.tsx\n→ getClinicBySlug → fetchClinicBySlug", "✅ GET /clinics/{slug}", "M"),
        ("P0-3", "Страница /cabinet/patient/appointments (404)",
         "Создать: frontend/src/app/cabinet/patient/appointments/page.tsx", "✅ GET /appointments/patient", "M"),
        ("P0-4", "Страница /cabinet/doctor/schedule (404)",
         "Создать: frontend/src/app/cabinet/doctor/schedule/page.tsx", "✅ GET /doctors/{id}/schedule", "M"),
        ("P0-5", "ClinicsSection на главной → реальные данные + правильные ссылки",
         "frontend/src/components/home/ClinicsSection.tsx\n→ fetchClinics() + href=/clinic/{slug}", "✅ GET /clinics", "S"),
        ("P0-6", "Медкарта: убрать чужие данные (хотя бы подставить реальное имя)",
         "frontend/src/app/cabinet/patient/medcard/page.tsx\n→ добавить GET /auth/me", "✅ GET /auth/me", "S"),
    ]
)

add_heading("P1 — Важные (следующий спринт)", level=2)
add_table(
    ["#", "Задача", "Файл", "Backend", "Сложность"],
    [
        ("P1-1", "Фильтры поиска ДМС/Метро/Онлайн",
         "lib/api.ts (apiDoctorToDoctor)\n+ Doctor-модель + миграция Alembic", "Нужно добавить поля", "M"),
        ("P1-2", "Избранные врачи (полная фича)",
         "Новая модель Favorite\n+ 3 endpoints\n+ страница /favorites\n+ кнопка ♡", "❌ Нужно всё", "L"),
        ("P1-3", "/cabinet/doctor/settings (404)",
         "Создать страницу + PATCH /doctors/{id}/profile", "❌ Нужен endpoint", "M"),
        ("P1-4", "/cabinet/clinic/settings — форма редактирования",
         "settings/page.tsx + PATCH /clinics/{id}", "❌ Нужен endpoint", "M"),
        ("P1-5", "Бонус-виджет на главной ЛК — реальная сумма",
         "cabinet/patient/page.tsx: client-side GET /auth/me", "✅", "S"),
        ("P1-6", "Бонус-виджет «Получить» — убрать или реализовать",
         "cabinet/patient/bonuses/page.tsx", "❌ Нужен endpoint", "M"),
    ]
)

add_heading("P2 — Улучшения (дальнейшая дорожная карта)", level=2)
add_table(
    ["Задача", "Что нужно", "Сложность"],
    [
        ("Медкарта — реальные данные", "Модель MedicalRecord + API + redesign страницы", "XL"),
        ("Семейный профиль — реальные данные", "Модель FamilyMember + API + страница", "XL"),
        ("PDF-экспорт медкарты", "puppeteer или pdfmake на бэкенде", "L"),
        ("Реальные отзывы на /doctor/[slug]", "Модель Review + API", "L"),
        ("Бонусы «Получить» — реальное погашение", "Модель BonusReward + endpoint + UI", "L"),
        ("Статьи на главной — из БД", "Модель Article + CMS-like API", "L"),
        ("StatsSection — реальные цифры", "Endpoint /stats/platform", "S"),
        ("DoctorsSection на главной — из API", "fetchDoctors() в HomePage", "S"),
        ("«Следующий приём» в ЛК — из записей", "GET /appointments/patient → nearest", "S"),
        ("Health stats в ЛК — убрать или из медкарты", "Убрать hardcoded секцию", "S"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. ТРЕБОВАНИЯ ДЛЯ PRODUCTION
# ══════════════════════════════════════════════════════════════

add_heading("8. Требования для перехода в Production (med-as.ru)", level=1)

add_heading("Блокеры — нельзя запускать без этого", level=2)
add_table(
    ["Требование", "Статус", "Приоритет"],
    [
        ("Профили 3 врачей (P0-1)", "❌ Не сделано", "CRITICAL"),
        ("Профили 3 клиник (P0-2)", "❌ Не сделано", "CRITICAL"),
        ("/cabinet/patient/appointments (P0-3)", "❌ Не сделано", "CRITICAL"),
        ("Клиники на главной (P0-5)", "❌ Не сделано", "CRITICAL"),
        ("Убрать чужие данные из медкарты (P0-6)", "❌ Не сделано", "HIGH"),
        ("Реальные Flash Call (websms.ru контракт)", "⚠️ PENDING", "CRITICAL"),
        ("/cabinet/doctor/schedule (P0-4)", "❌ Не сделано", "HIGH"),
        ("Фильтры поиска (P1-1)", "❌ Не сделано", "HIGH"),
    ]
)

add_heading("Минимальный путь пользователя БЕЗ ни одной 404", level=2)
add_para(
    "Для демонстрации инвестору или клиенту — только этот маршрут и только эти врачи:",
    italic=True
)
add_bullet("saas.med-as.ru/ → поиск (не кликать на клиники!) → /search?q=кардиолог")
add_bullet("/search → кликнуть ТОЛЬКО anna-sokolova, igor-petrov или maria-kozlova")
add_bullet("/doctor/anna-sokolova → выбрать дату → Записаться")
add_bullet("/doctor/anna-sokolova/booking → заполнить → Отправить")
add_bullet("/cabinet/patient → видны записи (НЕ нажимать 'Приёмы', 'Избранные'!)")
add_bullet("/cabinet/patient/bonuses → реальный баланс и история")

add_heading("Тестовые данные", level=2)
add_table(
    ["Роль", "Телефон", "Код", "Что тестировать"],
    [
        ("Пациент (новый)", "Любой номер", "[МАСТЕР-КОД]", "Регистрация, запись, ЛК"),
        ("Пациент (тест)", "+79271915291", "[МАСТЕР-КОД]", "Уже зарегистрирован, есть записи"),
        ("Клиника (тест)", "+70000000001", "123456", "Дашборд, врачи, расписание, отчёты"),
        ("Врач", "Нужен аккаунт с doctor_id", "[МАСТЕР-КОД]", "ЛК врача"),
    ]
)

# ──────────────────────────────────────────────
# ИТОГОВЫЕ РЕКОМЕНДАЦИИ
# ──────────────────────────────────────────────

add_heading("Рекомендации (итог аудита)", level=1)

rec = [
    ("R1 — КРИТИЧЕСКИ ВАЖНО", "Запустить 5 задач P0 до любого показа клиентам. "
     "Сейчас каждый стандартный клик пользователя заканчивается 404."),
    ("R2 — ДАННЫЕ ПАЦИЕНТОВ", "Медкарта и семья показывают чужие данные ('Алекс Стерлинг'). "
     "Любой пациент видит их как свои — это баг доверия, критичный для медицинского приложения."),
    ("R3 — ФИЛЬТРЫ ПОИСКА", "4 фильтра (ДМС, метро, онлайн, выезд) всегда дают 0 результатов. "
     "Решение: добавить поля в Doctor-модель + Alembic-миграция + маппинг в api.ts."),
    ("R4 — websms.ru", "Весь OTP-флоу сейчас работает только через мастер-пароль. "
     "Необходимо подписать контракт с websms.ru до публичного запуска."),
    ("R5 — СТАТИКА → API", "lib/doctors.ts (346 строк) и lib/clinics.ts (796 строк) — главная техническая долг. "
     "Переключение /doctor/[slug] и /clinic/[slug] на реальный API решит 6 из 16 broken flows сразу."),
    ("R6 — КЛИНИКА MVP", "ЛК клиники — самая зрелая часть системы. "
     "При продажах B2B демонстрировать именно её (аналитика + расписание + врачи)."),
    ("R7 — ВРАЧ BLOCKER", "Врач не может управлять расписанием и видеть настройки (404). "
     "Для онбординга первых реальных врачей P0-4 + P1-3 — обязательны."),
    ("R8 — БЕЗ МОК-ДАННЫХ", "Любые hardcoded данные в production создают риски доверия. "
     "Приоритет: убрать все mock до перехода на prod-домен med-as.ru."),
]
for title_r, text in rec:
    add_heading(title_r, level=4)
    add_para(text)

# ──────────────────────────────────────────────
# FOOTER
# ──────────────────────────────────────────────
doc.add_page_break()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = f.add_run(
    f"MEDAS Platform  |  Аудит на основе реального кода  |  "
    f"{datetime.date.today().strftime('%d.%m.%Y')}  |  Конфиденциально"
)
r.font.color.rgb = GRAY
r.font.size = Pt(9)

doc.save(DOCX_PATH)
print(f"✅ Saved: {DOCX_PATH}")
